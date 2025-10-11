"""
File processor utility for extracting text from various file formats.
Supports: TXT, PDF, DOCX
"""

import io
import logging
from pathlib import Path
from typing import Union, BinaryIO

# PDF processing
try:
    import PyPDF2
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# Word processing
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

logger = logging.getLogger(__name__)


class UnsupportedFileTypeError(Exception):
    """Raised when file type is not supported"""
    pass


def extract_text_from_txt(file_obj: BinaryIO) -> str:
    """
    Extract text from TXT file.

    Args:
        file_obj: File object

    Returns:
        Extracted text
    """
    try:
        content = file_obj.read()
        # Try UTF-8 first
        try:
            return content.decode('utf-8')
        except UnicodeDecodeError:
            # Fallback to latin-1
            return content.decode('latin-1', errors='ignore')
    except Exception as e:
        logger.error(f"Error reading TXT file: {e}")
        raise


def extract_text_from_pdf(file_obj: BinaryIO) -> str:
    """
    Extract text from PDF file using pdfplumber (better for resumes).
    Falls back to PyPDF2 if pdfplumber fails.

    Args:
        file_obj: File object

    Returns:
        Extracted text
    """
    if not PDF_AVAILABLE:
        raise UnsupportedFileTypeError("PDF processing libraries not installed. Install with: pip install PyPDF2 pdfplumber")

    text = ""

    try:
        # Method 1: pdfplumber (better for formatted documents like resumes)
        with pdfplumber.open(file_obj) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"

        if text.strip():
            return text.strip()

    except Exception as e:
        logger.warning(f"pdfplumber failed, trying PyPDF2: {e}")

        # Method 2: PyPDF2 (fallback)
        try:
            file_obj.seek(0)  # Reset file pointer
            pdf_reader = PyPDF2.PdfReader(file_obj)

            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"

            if text.strip():
                return text.strip()

        except Exception as e2:
            logger.error(f"PyPDF2 also failed: {e2}")
            raise Exception("Failed to extract text from PDF using both methods")

    if not text.strip():
        raise Exception("No text could be extracted from PDF")

    return text.strip()


def extract_text_from_docx(file_obj: BinaryIO) -> str:
    """
    Extract text from DOCX (Word) file.

    Args:
        file_obj: File object

    Returns:
        Extracted text
    """
    if not DOCX_AVAILABLE:
        raise UnsupportedFileTypeError("DOCX processing library not installed. Install with: pip install python-docx")

    try:
        doc = Document(file_obj)

        # Extract text from paragraphs
        text_parts = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)

        text = "\n".join(text_parts)

        if not text.strip():
            raise Exception("No text could be extracted from DOCX")

        return text.strip()

    except Exception as e:
        logger.error(f"Error reading DOCX file: {e}")
        raise


def extract_text_from_file(
    file_obj: Union[BinaryIO, io.BytesIO],
    filename: str = None
) -> str:
    """
    Extract text from file based on extension.
    Supports: .txt, .pdf, .docx

    Args:
        file_obj: File object or BytesIO
        filename: Original filename (to determine type)

    Returns:
        Extracted text

    Raises:
        UnsupportedFileTypeError: If file type is not supported
        Exception: If extraction fails
    """
    # Determine file type from filename
    if filename:
        ext = Path(filename).suffix.lower()
    else:
        raise ValueError("Filename required to determine file type")

    logger.info(f"Extracting text from {ext} file: {filename}")

    # Route to appropriate extractor
    if ext == '.txt':
        return extract_text_from_txt(file_obj)

    elif ext == '.pdf':
        return extract_text_from_pdf(file_obj)

    elif ext in ['.docx', '.doc']:
        if ext == '.doc':
            raise UnsupportedFileTypeError(
                "Old .doc format not supported. Please convert to .docx or save as PDF/TXT"
            )
        return extract_text_from_docx(file_obj)

    else:
        raise UnsupportedFileTypeError(
            f"File type '{ext}' not supported. Supported types: .txt, .pdf, .docx"
        )


def validate_file_size(file_obj: BinaryIO, max_size_mb: int = 10) -> bool:
    """
    Validate file size is within limits.

    Args:
        file_obj: File object
        max_size_mb: Maximum size in MB

    Returns:
        True if valid, raises exception otherwise
    """
    file_obj.seek(0, 2)  # Seek to end
    size_bytes = file_obj.tell()
    file_obj.seek(0)  # Reset to start

    size_mb = size_bytes / (1024 * 1024)

    if size_mb > max_size_mb:
        raise Exception(f"File too large: {size_mb:.2f}MB (max: {max_size_mb}MB)")

    return True


def get_file_info(file_obj: BinaryIO, filename: str) -> dict:
    """
    Get information about uploaded file.

    Args:
        file_obj: File object
        filename: Original filename

    Returns:
        Dictionary with file info
    """
    file_obj.seek(0, 2)
    size_bytes = file_obj.tell()
    file_obj.seek(0)

    return {
        "filename": filename,
        "extension": Path(filename).suffix.lower(),
        "size_bytes": size_bytes,
        "size_mb": round(size_bytes / (1024 * 1024), 2),
    }
