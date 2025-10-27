"""
Document generator for creating professional Word resumes from enhanced data.
"""

import io
from typing import Dict, Any, Optional
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class ResumeDocumentGenerator:
    """Generate professional Word documents from resume data"""

    def __init__(self):
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required. Install with: pip install python-docx")

        self.doc = Document()
        self._setup_styles()

    def _setup_styles(self):
        """Setup custom styles for the document"""
        styles = self.doc.styles

        # Header style (for name)
        if 'ResumeHeader' not in styles:
            header_style = styles.add_style('ResumeHeader', WD_STYLE_TYPE.PARAGRAPH)
            header_font = header_style.font
            header_font.name = 'Calibri'
            header_font.size = Pt(24)
            header_font.bold = True
            header_font.color.rgb = RGBColor(0, 51, 102)
            header_style.paragraph_format.space_after = Pt(6)

        # Section header style
        if 'SectionHeader' not in styles:
            section_style = styles.add_style('SectionHeader', WD_STYLE_TYPE.PARAGRAPH)
            section_font = section_style.font
            section_font.name = 'Calibri'
            section_font.size = Pt(14)
            section_font.bold = True
            section_font.color.rgb = RGBColor(0, 51, 102)
            section_style.paragraph_format.space_before = Pt(12)
            section_style.paragraph_format.space_after = Pt(6)

        # Job title style
        if 'JobTitle' not in styles:
            job_style = styles.add_style('JobTitle', WD_STYLE_TYPE.PARAGRAPH)
            job_font = job_style.font
            job_font.name = 'Calibri'
            job_font.size = Pt(11)
            job_font.bold = True
            job_style.paragraph_format.space_after = Pt(3)

    def _add_section_divider(self):
        """Add a horizontal line divider"""
        para = self.doc.add_paragraph()
        para.add_run('_' * 80)
        para.runs[0].font.color.rgb = RGBColor(200, 200, 200)

    def generate_from_enhanced_data(
        self,
        enhanced_data: Dict[str, Any],
        original_data: Dict[str, Any]
    ) -> io.BytesIO:
        """
        Generate Word document from enhanced resume data.

        Args:
            enhanced_data: Enhanced resume data from AI (can be formatted_resume from template formatter)
            original_data: Original parsed resume data

        Returns:
            BytesIO object containing the Word document
        """
        # Check if this is template-formatted data (has 'sections' array)
        if 'sections' in enhanced_data:
            return self._generate_from_formatted_sections(enhanced_data, original_data)

        # Otherwise use the original enhancement workflow
        # Use enhanced data if available, fall back to original
        resume_data = self._merge_enhanced_data(enhanced_data, original_data)

        # Add header (name and contact)
        self._add_header(resume_data.get('personal_info', {}))

        # Add summary
        if resume_data.get('summary'):
            self._add_summary(resume_data['summary'])

        # Add work experience
        if resume_data.get('work_experience'):
            self._add_work_experience(resume_data['work_experience'], enhanced_data)

        # Add skills
        if resume_data.get('skills'):
            self._add_skills(resume_data['skills'])

        # Add education
        if resume_data.get('education'):
            self._add_education(resume_data['education'])

        # Add enhancement footer
        self._add_enhancement_footer(enhanced_data)

        # Save to BytesIO
        doc_io = io.BytesIO()
        self.doc.save(doc_io)
        doc_io.seek(0)
        return doc_io

    def _generate_from_formatted_sections(
        self,
        formatted_data: Dict[str, Any],
        original_data: Dict[str, Any]
    ) -> io.BytesIO:
        """
        Generate document from template-formatted resume with sections in specific order.

        Args:
            formatted_data: Formatted resume with 'sections' array
            original_data: Original parsed resume (fallback)

        Returns:
            BytesIO containing the Word document
        """
        # Add header (name and contact)
        personal_info = formatted_data.get('personal_info', original_data.get('personal_info', {}))
        self._add_header(personal_info)

        # Process sections in the order specified by template formatter
        sections = formatted_data.get('sections', [])

        for section in sections:
            section_name = section.get('name', '').lower()
            section_content = section.get('content', {})

            # Map section names to appropriate handlers
            if 'summary' in section_name or 'objective' in section_name or 'profile' in section_name:
                if section_content:
                    self._add_summary_from_formatted(section)
            elif 'experience' in section_name or 'work' in section_name or 'employment' in section_name:
                if section_content:
                    self._add_experience_from_formatted(section)
            elif 'skill' in section_name:
                if section_content:
                    self._add_skills_from_formatted(section)
            elif 'education' in section_name:
                if section_content:
                    self._add_education_from_formatted(section)
            elif 'certification' in section_name or 'certificate' in section_name:
                if section_content:
                    self._add_generic_section(section)
            elif 'project' in section_name:
                if section_content:
                    self._add_generic_section(section)
            else:
                # Generic section handler for any other sections
                if section_content:
                    self._add_generic_section(section)

        # Save to BytesIO
        doc_io = io.BytesIO()
        self.doc.save(doc_io)
        doc_io.seek(0)
        return doc_io

    def _add_summary_from_formatted(self, section: Dict):
        """Add summary section from formatted data"""
        section_name = section.get('name', 'PROFESSIONAL SUMMARY')
        content = section.get('content', {})

        self.doc.add_paragraph(section_name.upper(), style='SectionHeader')
        self._add_section_divider()

        # Handle different summary formats
        if isinstance(content, str):
            summary_para = self.doc.add_paragraph(content)
        elif isinstance(content, dict):
            summary_text = content.get('summary_text', content.get('text', ''))
            summary_para = self.doc.add_paragraph(summary_text)
        else:
            return

        summary_para.paragraph_format.space_after = Pt(12)

    def _add_experience_from_formatted(self, section: Dict):
        """Add work experience section from formatted data"""
        section_name = section.get('name', 'WORK EXPERIENCE')
        content = section.get('content', {})

        self.doc.add_paragraph(section_name.upper(), style='SectionHeader')
        self._add_section_divider()

        # Handle list of experiences
        experiences = content if isinstance(content, list) else content.get('experiences', [])

        for exp in experiences:
            if not isinstance(exp, dict):
                continue

            # Job title and company
            job_para = self.doc.add_paragraph(style='JobTitle')
            title = exp.get('title', exp.get('position', 'Position'))
            company = exp.get('company', 'Company')
            job_para.add_run(f"{title} - {company}")

            # Location if available
            if exp.get('location'):
                job_para.add_run(f" | {exp['location']}")

            # Dates
            start_date = exp.get('start_date', 'Start')
            end_date = exp.get('end_date', 'Present')
            date_para = self.doc.add_paragraph()
            date_run = date_para.add_run(f"{start_date} - {end_date}")
            date_run.font.italic = True
            date_run.font.size = Pt(10)

            # Achievements/responsibilities
            achievements = exp.get('achievements', exp.get('responsibilities', []))
            for achievement in achievements:
                ach_para = self.doc.add_paragraph(achievement, style='List Bullet')
                ach_para.paragraph_format.left_indent = Inches(0.25)

            # Technologies if available
            if exp.get('technologies'):
                tech_para = self.doc.add_paragraph()
                tech_para.add_run('Technologies: ').bold = True
                tech_para.add_run(', '.join(exp['technologies']))
                tech_para.runs[-1].font.italic = True
                tech_para.runs[-1].font.size = Pt(10)

            self.doc.add_paragraph()  # Spacer between jobs

    def _add_skills_from_formatted(self, section: Dict):
        """Add skills section from formatted data"""
        section_name = section.get('name', 'SKILLS')
        content = section.get('content', {})

        self.doc.add_paragraph(section_name.upper(), style='SectionHeader')
        self._add_section_divider()

        # Handle different skill formats
        if isinstance(content, dict):
            # Technical skills
            if content.get('technical'):
                tech_para = self.doc.add_paragraph()
                tech_para.add_run('Technical: ').bold = True
                skills_list = content['technical'] if isinstance(content['technical'], list) else [content['technical']]
                tech_para.add_run(', '.join(skills_list))

            # Soft skills
            if content.get('soft_skills'):
                soft_para = self.doc.add_paragraph()
                soft_para.add_run('Soft Skills: ').bold = True
                skills_list = content['soft_skills'] if isinstance(content['soft_skills'], list) else [content['soft_skills']]
                soft_para.add_run(', '.join(skills_list))

            # Domain knowledge
            if content.get('domain_knowledge'):
                domain_para = self.doc.add_paragraph()
                domain_para.add_run('Domain Knowledge: ').bold = True
                skills_list = content['domain_knowledge'] if isinstance(content['domain_knowledge'], list) else [content['domain_knowledge']]
                domain_para.add_run(', '.join(skills_list))

            # Tools
            if content.get('tools'):
                tools_para = self.doc.add_paragraph()
                tools_para.add_run('Tools: ').bold = True
                skills_list = content['tools'] if isinstance(content['tools'], list) else [content['tools']]
                tools_para.add_run(', '.join(skills_list))
        elif isinstance(content, list):
            # Simple list of skills
            skills_para = self.doc.add_paragraph(', '.join(content))
        elif isinstance(content, str):
            # String of skills
            skills_para = self.doc.add_paragraph(content)

        self.doc.add_paragraph()  # Spacer

    def _add_education_from_formatted(self, section: Dict):
        """Add education section from formatted data"""
        section_name = section.get('name', 'EDUCATION')
        content = section.get('content', {})

        self.doc.add_paragraph(section_name.upper(), style='SectionHeader')
        self._add_section_divider()

        # Handle list of education entries
        education = content if isinstance(content, list) else content.get('education', [])

        for edu in education:
            if not isinstance(edu, dict):
                continue

            # Degree and field
            degree_para = self.doc.add_paragraph(style='JobTitle')
            degree = edu.get('degree', 'Degree')
            field = edu.get('field', '')
            if field:
                degree_text = f"{degree} in {field}"
            else:
                degree_text = degree
            degree_para.add_run(degree_text)

            # Institution and date
            institution = edu.get('institution', 'Institution')
            grad_date = edu.get('graduation_date', edu.get('end_date', 'Year'))
            inst_para = self.doc.add_paragraph()
            inst_para.add_run(f"{institution} | {grad_date}")
            inst_para.runs[0].font.italic = True

            # GPA if available
            if edu.get('gpa'):
                gpa_para = self.doc.add_paragraph()
                gpa_para.add_run(f"GPA: {edu['gpa']}")
                gpa_para.runs[0].font.size = Pt(10)

            self.doc.add_paragraph()  # Spacer

    def _add_generic_section(self, section: Dict):
        """Add any generic section from formatted data"""
        section_name = section.get('name', 'SECTION')
        content = section.get('content', {})

        self.doc.add_paragraph(section_name.upper(), style='SectionHeader')
        self._add_section_divider()

        # Handle different content types
        if isinstance(content, str):
            self.doc.add_paragraph(content)
        elif isinstance(content, list):
            for item in content:
                if isinstance(item, str):
                    self.doc.add_paragraph(item, style='List Bullet')
                elif isinstance(item, dict):
                    # Add item title if exists
                    if item.get('title'):
                        title_para = self.doc.add_paragraph(style='JobTitle')
                        title_para.add_run(item['title'])
                    # Add item description
                    if item.get('description'):
                        self.doc.add_paragraph(item['description'])
        elif isinstance(content, dict):
            for key, value in content.items():
                para = self.doc.add_paragraph()
                para.add_run(f"{key}: ").bold = True
                if isinstance(value, list):
                    para.add_run(', '.join(str(v) for v in value))
                else:
                    para.add_run(str(value))

        self.doc.add_paragraph()  # Spacer

    def _merge_enhanced_data(self, enhanced: Dict, original: Dict) -> Dict:
        """Merge enhanced data with original, prioritizing enhanced content"""
        result = original.copy()

        # Merge enhanced summary
        if enhanced.get('enhanced_summary'):
            if 'summary' not in result:
                result['summary'] = {}
            result['summary']['summary_text'] = enhanced['enhanced_summary']

        # Merge enhanced experience
        if enhanced.get('enhanced_experience'):
            for i, enh_exp in enumerate(enhanced['enhanced_experience']):
                if i < len(result.get('work_experience', [])):
                    # Replace achievements with enhanced ones
                    if enh_exp.get('achievements'):
                        enhanced_achievements = [
                            ach.get('enhanced', ach.get('original', ''))
                            for ach in enh_exp['achievements']
                        ]
                        result['work_experience'][i]['achievements'] = enhanced_achievements

        return result

    def _add_header(self, personal_info: Dict):
        """Add name and contact information"""
        # Name
        name_para = self.doc.add_paragraph(personal_info.get('full_name', 'Name'), style='ResumeHeader')
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Contact info
        contact_parts = []
        if personal_info.get('email'):
            contact_parts.append(personal_info['email'])
        if personal_info.get('phone'):
            contact_parts.append(personal_info['phone'])
        if personal_info.get('location'):
            contact_parts.append(personal_info['location'])

        if contact_parts:
            contact_para = self.doc.add_paragraph(' | '.join(contact_parts))
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_para.runs[0].font.size = Pt(10)

        if personal_info.get('linkedin_url'):
            linkedin_para = self.doc.add_paragraph(personal_info['linkedin_url'])
            linkedin_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            linkedin_para.runs[0].font.size = Pt(10)
            linkedin_para.runs[0].font.color.rgb = RGBColor(0, 102, 204)

        self.doc.add_paragraph()  # Spacer

    def _add_summary(self, summary: Dict):
        """Add professional summary section"""
        self.doc.add_paragraph('PROFESSIONAL SUMMARY', style='SectionHeader')
        self._add_section_divider()

        summary_para = self.doc.add_paragraph(summary.get('summary_text', ''))
        summary_para.paragraph_format.space_after = Pt(12)

    def _add_work_experience(self, work_experience: list, enhanced_data: Dict):
        """Add work experience section"""
        self.doc.add_paragraph('WORK EXPERIENCE', style='SectionHeader')
        self._add_section_divider()

        for exp in work_experience:
            # Job title and company
            job_para = self.doc.add_paragraph(style='JobTitle')
            job_para.add_run(f"{exp.get('title', 'Position')} - {exp.get('company', 'Company')}")

            # Dates
            date_para = self.doc.add_paragraph()
            date_run = date_para.add_run(f"{exp.get('start_date', 'Start')} - {exp.get('end_date', 'Present')}")
            date_run.font.italic = True
            date_run.font.size = Pt(10)

            # Achievements
            for achievement in exp.get('achievements', []):
                ach_para = self.doc.add_paragraph(achievement, style='List Bullet')
                ach_para.paragraph_format.left_indent = Inches(0.25)

            # Technologies
            if exp.get('technologies'):
                tech_para = self.doc.add_paragraph()
                tech_para.add_run('Technologies: ').bold = True
                tech_para.add_run(', '.join(exp['technologies']))
                tech_para.runs[-1].font.italic = True
                tech_para.runs[-1].font.size = Pt(10)

            self.doc.add_paragraph()  # Spacer between jobs

    def _add_skills(self, skills: Dict):
        """Add skills section"""
        self.doc.add_paragraph('SKILLS', style='SectionHeader')
        self._add_section_divider()

        # Technical skills
        if skills.get('technical'):
            tech_para = self.doc.add_paragraph()
            tech_para.add_run('Technical: ').bold = True
            tech_para.add_run(', '.join(skills['technical']))

        # Soft skills
        if skills.get('soft_skills'):
            soft_para = self.doc.add_paragraph()
            soft_para.add_run('Soft Skills: ').bold = True
            soft_para.add_run(', '.join(skills['soft_skills']))

        # Domain knowledge
        if skills.get('domain_knowledge'):
            domain_para = self.doc.add_paragraph()
            domain_para.add_run('Domain Knowledge: ').bold = True
            domain_para.add_run(', '.join(skills['domain_knowledge']))

        # Tools
        if skills.get('tools'):
            tools_para = self.doc.add_paragraph()
            tools_para.add_run('Tools: ').bold = True
            tools_para.add_run(', '.join(skills['tools']))

        self.doc.add_paragraph()  # Spacer

    def _add_education(self, education: list):
        """Add education section"""
        self.doc.add_paragraph('EDUCATION', style='SectionHeader')
        self._add_section_divider()

        for edu in education:
            # Degree and field
            degree_para = self.doc.add_paragraph(style='JobTitle')
            degree_text = f"{edu.get('degree', 'Degree')} in {edu.get('field', 'Field')}"
            degree_para.add_run(degree_text)

            # Institution and date
            inst_para = self.doc.add_paragraph()
            inst_para.add_run(f"{edu.get('institution', 'Institution')} | {edu.get('graduation_date', 'Year')}")
            inst_para.runs[0].font.italic = True

            self.doc.add_paragraph()  # Spacer

    def _add_enhancement_footer(self, enhanced_data: Dict):
        """Add footer showing enhancement information"""
        if not enhanced_data:
            return

        self.doc.add_page_break()
        self.doc.add_paragraph('ENHANCEMENT DETAILS', style='SectionHeader')
        self._add_section_divider()

        # ATS Score
        if enhanced_data.get('ats_score'):
            ats = enhanced_data['ats_score']
            ats_para = self.doc.add_paragraph()
            ats_para.add_run('ATS Score Improvement: ').bold = True
            ats_para.add_run(f"{ats.get('before', 0)}% → {ats.get('after', 0)}% ")
            ats_para.add_run(f"(+{ats.get('after', 0) - ats.get('before', 0)}%)")
            ats_para.runs[-1].font.color.rgb = RGBColor(0, 153, 0)

        # Keywords added
        if enhanced_data.get('keywords_added'):
            kw_para = self.doc.add_paragraph()
            kw_para.add_run('Keywords Added: ').bold = True
            kw_para.add_run(', '.join(enhanced_data['keywords_added']))

        # Changes summary
        if enhanced_data.get('change_summary'):
            self.doc.add_paragraph()
            self.doc.add_paragraph('Changes Made:', style='JobTitle')
            for change in enhanced_data['change_summary']:
                self.doc.add_paragraph(change, style='List Bullet')

        # Generation timestamp
        self.doc.add_paragraph()
        timestamp_para = self.doc.add_paragraph()
        timestamp_para.add_run(f'Generated: {datetime.now().strftime("%B %d, %Y at %I:%M %p")}')
        timestamp_para.runs[0].font.size = Pt(8)
        timestamp_para.runs[0].font.italic = True
        timestamp_para.runs[0].font.color.rgb = RGBColor(128, 128, 128)
        timestamp_para.alignment = WD_ALIGN_PARAGRAPH.CENTER


def generate_enhanced_resume_docx(
    enhanced_data: Dict[str, Any],
    original_data: Dict[str, Any]
) -> io.BytesIO:
    """
    Generate a Word document with enhanced resume.

    Args:
        enhanced_data: Enhanced resume data from enhancement agent
        original_data: Original parsed resume data

    Returns:
        BytesIO containing the Word document
    """
    generator = ResumeDocumentGenerator()
    return generator.generate_from_enhanced_data(enhanced_data, original_data)
